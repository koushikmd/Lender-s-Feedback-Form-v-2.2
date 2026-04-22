"""DOCX generator for Lender's Feedback Form - runs in Pyodide."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


def _set_cell_bold(cell, bold=True):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold


def _make_cell_bold(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)


def _shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def generate_docx(data: dict) -> bytes:
    doc = Document()
    # Set default margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Lender's Feedback Form for Credit Rating of Clients")
    title_run.bold = True
    title_run.underline = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()

    # Section: General Information
    _add_section_header(doc, "General Information")
    gen_data = [
        ("Business Name", data.get("business_name", "")),
        ("Address", data.get("address", "")),
        ("Key Person", data.get("key_person", "")),
        ("Contact Number", data.get("contact_number", "")),
        ("Ownership Type", data.get("ownership_type", "")),
        ("RM Name", data.get("rm_name", "")),
        ("RM Contact Number", ""),  # Always blank - user fills manually
    ]
    t = doc.add_table(rows=len(gen_data), cols=2)
    t.style = 'Table Grid'
    for i, (label, value) in enumerate(gen_data):
        _make_cell_bold(t.rows[i].cells[0], label)
        _make_cell_bold(t.rows[i].cells[1], str(value) if value else "")
        t.rows[i].cells[0].width = Inches(2.2)
        t.rows[i].cells[1].width = Inches(4.3)

    doc.add_paragraph()

    # Section: Business Dynamics
    _add_section_header(doc, "Business Dynamics")
    bd_data = [
        ("Background", data.get("background", "")),
        ("Major Suppliers", data.get("major_suppliers", "")),
        ("Major Clients", data.get("major_clients", "")),
    ]
    t = doc.add_table(rows=len(bd_data), cols=2)
    t.style = 'Table Grid'
    for i, (label, value) in enumerate(bd_data):
        _make_cell_bold(t.rows[i].cells[0], label)
        # Multi-line content
        cell = t.rows[i].cells[1]
        cell.text = ""
        value_str = str(value) if value else ""
        paras = value_str.split("\n")
        for pi, para_text in enumerate(paras):
            if pi == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            r = p.add_run(para_text)
            r.bold = True
            r.font.size = Pt(11)
        t.rows[i].cells[0].width = Inches(2.2)
        t.rows[i].cells[1].width = Inches(4.3)

    doc.add_page_break()

    # Section: Financial Information
    _add_section_header(doc, "Financial Information")
    fin_data = [
        ("Sales/Revenue", data.get("sales_revenue", "")),
        ("Inventory", data.get("inventory", "")),
        ("A/C Receivable", data.get("ar", "")),
        ("A/C Payable", data.get("ap", "")),
        ("Gross Profit (%)", data.get("gp_pct", "")),
        ("Net Profit (%)", data.get("np_pct", "")),
    ]
    t = doc.add_table(rows=len(fin_data), cols=2)
    t.style = 'Table Grid'
    for i, (label, value) in enumerate(fin_data):
        _make_cell_bold(t.rows[i].cells[0], label)
        _make_cell_bold(t.rows[i].cells[1], str(value) if value else "")
        t.rows[i].cells[0].width = Inches(2.2)
        t.rows[i].cells[1].width = Inches(4.3)

    doc.add_paragraph()

    # Section: Long Term Liability
    _add_section_header(doc, "Long Term Liability")
    lt_list = data.get("long_term_liabilities") or []
    lt_headers = ["Bank Name", "Facility Type", "Limit",
                  f"Outstanding (as of {data.get('date', '')})",
                  "EMI", "Tenure", "Repayment Status"]
    t = doc.add_table(rows=len(lt_list) + 1, cols=len(lt_headers))
    t.style = 'Table Grid'
    for ci, h in enumerate(lt_headers):
        _make_cell_bold(t.rows[0].cells[ci], h)
        _shade_cell(t.rows[0].cells[ci], "D9E2F3")
    for ri, entry in enumerate(lt_list, start=1):
        row = t.rows[ri]
        _make_cell_bold(row.cells[0], entry.get("bank_name", ""))
        _make_cell_bold(row.cells[1], entry.get("facility_type", ""))
        _make_cell_bold(row.cells[2], entry.get("limit", ""))
        _make_cell_bold(row.cells[3], entry.get("outstanding", ""))
        _make_cell_bold(row.cells[4], entry.get("emi", ""))
        _make_cell_bold(row.cells[5], entry.get("tenure", ""))
        _make_cell_bold(row.cells[6], entry.get("repayment_status", "Regular"))

    doc.add_paragraph()

    # Section: Short Term Liability
    _add_section_header(doc, "Short Term Liability")
    st_list = data.get("short_term_liabilities") or []
    st_headers = ["Bank Name", "Facility Type", "Limit",
                  f"Outstanding (as of {data.get('date', '')})",
                  "Recycle Times", "Expiry Date", "Repayment Status"]
    t = doc.add_table(rows=len(st_list) + 1, cols=len(st_headers))
    t.style = 'Table Grid'
    for ci, h in enumerate(st_headers):
        _make_cell_bold(t.rows[0].cells[ci], h)
        _shade_cell(t.rows[0].cells[ci], "D9E2F3")
    for ri, entry in enumerate(st_list, start=1):
        row = t.rows[ri]
        _make_cell_bold(row.cells[0], entry.get("bank_name", ""))
        _make_cell_bold(row.cells[1], entry.get("facility_type", ""))
        _make_cell_bold(row.cells[2], entry.get("limit", ""))
        _make_cell_bold(row.cells[3], entry.get("outstanding", ""))
        _make_cell_bold(row.cells[4], entry.get("recycle_times", ""))
        _make_cell_bold(row.cells[5], entry.get("expiry_date", ""))
        _make_cell_bold(row.cells[6], entry.get("repayment_status", "Regular"))

    doc.add_paragraph()

    # Import/Export Line
    p = doc.add_paragraph()
    r = p.add_run(f"Import/Export: {data.get('import_export', 'N/A')}")
    r.bold = True
    r.font.size = Pt(11)

    # Credit Summation Line
    p = doc.add_paragraph()
    cs = data.get('credit_summation', '')
    r = p.add_run(f"Credit Summation: BDT {cs}" if cs else "Credit Summation: ")
    r.bold = True
    r.font.size = Pt(11)

    # Save to bytes buffer
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.underline = True
    r.font.size = Pt(12)
